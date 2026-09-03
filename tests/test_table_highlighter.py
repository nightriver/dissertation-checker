"""Перевірки безпечного підсвічування порівняльних таблиць DOCX."""

from __future__ import annotations

import io
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Inches, Pt, RGBColor
import pytest

from table_highlighter.layout import logical_rows
from table_highlighter.matcher import align, left_statuses, right_statuses
from table_highlighter.processor import DocumentValidationError, inspect_tables, process_document
from table_highlighter.types import HighlightOptions
from table_highlighter.zones import annotate_left_cell, annotate_right_cell, paragraph_text


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Створює реальне w:hyperlink, а не текст URL для перевірки збереження."""
    part = paragraph.part
    relation = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _table_docx(rows: list[tuple[list[str], list[str]]], *, link_in_right: bool = False) -> bytes:
    document = Document()
    table = document.add_table(rows=0, cols=2)
    for row_index, (left, right) in enumerate(rows):
        row = table.add_row()
        for cell, paragraphs in zip(row.cells, (left, right)):
            cell.text = paragraphs[0]
            for text in paragraphs[1:]:
                cell.add_paragraph(text)
        if link_in_right and row_index == 0:
            _add_hyperlink(row.cells[1].paragraphs[0], " активне посилання", "https://example.test/source")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _cell_texts(cell) -> list[str]:
    return [paragraph_text(paragraph) for paragraph in cell.paragraphs]


def _has_hyperlink(cell) -> bool:
    return bool(cell._tc.findall(".//" + qn("w:hyperlink")))


class TestZones:
    def test_right_comment_keeps_every_following_paragraph_plain(self):
        data = _table_docx([
            (["С. 1", "Текст", "", "Примітка"], ["Джерело", "С. 2", "Текст", "", "Посилання на:", "29. Джерело", "", "60. Інше джерело"]),
        ])
        document = Document(io.BytesIO(data))
        right = annotate_right_cell(document.tables[0].rows[0].cells[1])

        assert [item.zone for item in right.paragraphs] == ["plain", "plain", "text", "text", "plain", "plain", "plain", "plain"]
        assert [paragraph_text(item.paragraph) for item in right.paragraphs[-3:]] == ["29. Джерело", "", "60. Інше джерело"]

    def test_left_marks_text_only_before_first_empty_paragraph(self):
        data = _table_docx([(["С. 1", "Текст", "", "Коментар"], ["С. 2", "Текст"])])
        document = Document(io.BytesIO(data))
        left = annotate_left_cell(document.tables[0].rows[0].cells[0])
        assert [item.zone for item in left.paragraphs] == ["plain", "text", "plain", "plain"]


class TestMatcher:
    def test_threshold_100_disables_relaxed_short_word_match(self):
        data = _table_docx([(["С. 1", "суд"], ["С. 2", "суду"])])
        document = Document(io.BytesIO(data))
        left = annotate_left_cell(document.tables[0].rows[0].cells[0])
        right = annotate_right_cell(document.tables[0].rows[0].cells[1])
        result = align(left.text_paragraphs, right.text_paragraphs, 100, True)
        assert set(left_statuses(result)[1]) == {"diff"}
        assert set(right_statuses(result)[1]) == {"diff"}

    def test_normalized_hyphen_is_only_a_matching_key(self):
        data = _table_docx([(["С. 1", "Ін- тернет"], ["С. 2", "Інтернет"])])
        document = Document(io.BytesIO(data))
        left = annotate_left_cell(document.tables[0].rows[0].cells[0])
        right = annotate_right_cell(document.tables[0].rows[0].cells[1])
        result = align(left.text_paragraphs, right.text_paragraphs, 75, True)
        assert set(left_statuses(result)[1]) == {"match"}
        assert set(right_statuses(result)[1]) == {"match"}

    def test_punctuation_between_matches_does_not_inherit_highlight(self):
        data = _table_docx([(["С. 1", "слово — слово"], ["С. 2", "слово — слово"])])
        document = Document(io.BytesIO(data))
        left = annotate_left_cell(document.tables[0].rows[0].cells[0])
        right = annotate_right_cell(document.tables[0].rows[0].cells[1])
        result = align(left.text_paragraphs, right.text_paragraphs, 75, True)
        statuses = left_statuses(result)[1]
        assert statuses[6] is None
        assert statuses[0] == "match"
        assert statuses[-1] == "match"


class TestProcessor:
    def test_preserves_comment_tail_hyperlink_and_original_hyphen(self):
        data = _table_docx([
            (
                ["С. 1", "Ін- тернет суду", "", "Коментар"],
                ["Джерело", "С. 2", "Інтернет суд", "", "Посилання на:", "29. Джерело", "", "60. Інше джерело"],
            ),
        ], link_in_right=True)
        result = process_document(data, HighlightOptions())
        document = Document(io.BytesIO(result.document_bytes))
        record = logical_rows(document.tables[0])[0]
        left, right = record.row.cells

        assert result.stats.processed_rows == 1
        assert "Ін- тернет суду" in _cell_texts(left)
        assert _cell_texts(right)[-3:] == ["29. Джерело", "", "60. Інше джерело"]
        assert _has_hyperlink(record.header.cells[1])
        colors = [str(run.font.highlight_color) for paragraph in left.paragraphs for run in paragraph.runs]
        assert "YELLOW (7)" in colors

    def test_preserves_hyperlink_inside_compared_text(self):
        data = _table_docx([(["С. 1", "Спільний текст"], ["С. 2", "Спільний текст"])])
        document = Document(io.BytesIO(data))
        right = document.tables[0].rows[0].cells[1]
        _add_hyperlink(right.paragraphs[1], " посилання", "https://example.test/inside")
        raw = io.BytesIO()
        document.save(raw)

        result = process_document(raw.getvalue(), HighlightOptions())
        output = Document(io.BytesIO(result.document_bytes))
        cell = output.tables[0].rows[0].cells[1]
        assert _has_hyperlink(cell)
        assert "посилання" in paragraph_text(cell.paragraphs[1])

    def test_alignment_gives_markers_a_shared_row_below_bibliography(self):
        data = _table_docx([(["С. 1", "Однаковий текст"], ["Бібліографія", "С. 2", "Однаковий текст"])])
        result = process_document(data, HighlightOptions())
        document = Document(io.BytesIO(result.document_bytes))
        record = logical_rows(document.tables[0])[0]
        assert result.stats.aligned_rows == 1
        assert _cell_texts(record.header.cells[0]) == [""]
        assert _cell_texts(record.header.cells[1]) == ["Бібліографія"]
        assert [cell.paragraphs[0].text for cell in record.row.cells] == ["С. 1", "С. 2"]

    def test_reprocessing_does_not_accumulate_alignment_padding(self):
        data = _table_docx([(["С. 1", "Текст"], ["Джерело", "С. 2", "Текст"])])
        first = process_document(data, HighlightOptions())
        second = process_document(first.document_bytes, HighlightOptions())
        document = Document(io.BytesIO(second.document_bytes))
        record = logical_rows(document.tables[0])[0]
        assert len(document.tables[0].rows) == 2
        assert _cell_texts(record.header.cells[0]) == [""]
        assert _cell_texts(record.row.cells[0]) == ["С. 1", "Текст"]
        assert inspect_tables(first.document_bytes) == inspect_tables(second.document_bytes) == inspect_tables(data)
        assert second.stats.aligned_rows == 1

    def test_missing_marker_keeps_row_without_destroying_empty_paragraphs(self):
        data = _table_docx([(["Текст", "", "Інший текст"], ["С. 2", "Текст"])])
        result = process_document(data, HighlightOptions())
        document = Document(io.BytesIO(result.document_bytes))
        assert result.stats.skipped_rows == 1
        assert _cell_texts(document.tables[0].rows[0].cells[0]) == ["Текст", "", "Інший текст"]
        assert result.warnings[0].row_number == 1

    def test_inspect_tables_and_selected_table_leave_other_tables_unchanged(self):
        document = Document()
        first = document.add_table(rows=1, cols=2)
        first.cell(0, 0).text = "Не обробляти"
        first.cell(0, 1).text = "Не обробляти"
        second = document.add_table(rows=1, cols=2)
        second.cell(0, 0).text = "С. 1"
        second.cell(0, 0).add_paragraph("Текст")
        second.cell(0, 1).text = "С. 2"
        second.cell(0, 1).add_paragraph("Текст")
        raw = io.BytesIO()
        document.save(raw)

        assert inspect_tables(raw.getvalue()) == (
            type(inspect_tables(raw.getvalue())[0])(0, 1, 1),
            type(inspect_tables(raw.getvalue())[1])(1, 1, 1),
        )
        result = process_document(raw.getvalue(), HighlightOptions(table_index=1))
        output = Document(io.BytesIO(result.document_bytes))
        assert output.tables[0].cell(0, 0).text == "Не обробляти"
        assert result.stats.processed_rows == 1

    def test_rejects_document_without_table(self):
        document = Document()
        document.add_paragraph("Текст")
        raw = io.BytesIO()
        document.save(raw)
        with pytest.raises(DocumentValidationError, match="не містить таблиць"):
            process_document(raw.getvalue(), HighlightOptions())


class TestUniformFormatting:
    @pytest.mark.parametrize("part_name,child_name", [("footnotes", "footnote"), ("endnotes", "endnote"), ("comments", "comment")])
    def test_normalizes_note_parts_without_removing_their_content(self, part_name, child_name):
        document = Document(io.BytesIO(_table_docx([(["С. 1", "Текст"], ["С. 2", "Текст"])])))
        xml = (
            f'<w:{part_name} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:{child_name} w:id="1"><w:p><w:r><w:rPr><w:b/><w:i/>'
            '<w:color w:val="FF0000"/><w:sz w:val="40"/></w:rPr><w:t>Note</w:t></w:r>'
            f'</w:p></w:{child_name}></w:{part_name}>'
        ).encode()
        part = Part(PackURI(f"/word/{part_name}.xml"), f"application/vnd.openxmlformats-officedocument.wordprocessingml.{part_name}+xml", xml, document.part.package)
        document.part.relate_to(part, f"http://schemas.openxmlformats.org/officeDocument/2006/relationships/{part_name}")
        raw = io.BytesIO()
        document.save(raw)
        result = process_document(raw.getvalue(), HighlightOptions())
        with ZipFile(io.BytesIO(result.document_bytes)) as archive:
            note = parse_xml(archive.read(f"word/{part_name}.xml"))
        assert next(note.iter(qn("w:t"))).text == "Note"
        assert next(note.iter(qn("w:b"))).get(qn("w:val")) == "0"
        assert next(note.iter(qn("w:i"))).get(qn("w:val")) == "0"
        assert next(note.iter(qn("w:color"))).get(qn("w:val")) == "000000"
        assert next(note.iter(qn("w:sz"))).get(qn("w:val")) == "28"

    def test_resets_direct_and_inherited_formatting_everywhere_and_preserves_urls(self):
        document = Document(io.BytesIO(_table_docx([
            (["С. 1", "Спільний текст", "", "Пояснення"], ["Джерело", "С. 2", "Спільний текст"]),
            (["Без маркера"], ["Пропущений рядок"]),
        ], link_in_right=True)))
        document.add_paragraph("Заголовок поза таблицею", "Heading 1")
        other = document.add_table(rows=1, cols=1)
        other.cell(0, 0).text = "Інша таблиця"
        nested = other.cell(0, 0).add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "Вкладена таблиця"
        document.sections[0].header.paragraphs[0].text = "Колонтитул"
        document.sections[0].footer.paragraphs[0].text = "Нижній колонтитул"
        style = document.styles.add_style("MixedCharacterStyle", WD_STYLE_TYPE.CHARACTER)
        for font in (style.font, document.styles["Normal"].font, document.styles["Heading 1"].font):
            font.name = "Courier New"
            font.size = Pt(24)
            font.bold = True
            font.italic = True
            font.underline = True
            font.color.rgb = RGBColor(255, 0, 0)
            font.highlight_color = WD_COLOR_INDEX.GREEN
        source_parts = (document._element, document.sections[0].header._element, document.sections[0].footer._element)
        for root in source_parts:
            for run in root.iter(qn("w:r")):
                rpr = run.get_or_add_rPr()
                for name, value in (("b", "1"), ("i", "1"), ("u", "double"), ("color", "FF0000"), ("highlight", "green"), ("rStyle", style.style_id)):
                    item = OxmlElement(f"w:{name}")
                    item.set(qn("w:val"), value)
                    rpr.append(item)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "FF0000")
                rpr.append(shading)
        raw = io.BytesIO()
        document.save(raw)
        result = process_document(raw.getvalue(), HighlightOptions(font_name="Arial", font_size=11))
        output = Document(io.BytesIO(result.document_bytes))
        assert result.stats.processed_rows == 1 and result.stats.skipped_rows == 1
        for part in output.part.package.parts:
            root = getattr(part, "_element", None)
            if root is None:
                continue
            assert not list(root.iter(qn("w:shd")))
            for properties in root.iter(qn("w:rPr")):
                assert properties.find(qn("w:rStyle")) is None
                assert properties.find(qn("w:rFonts")).get(qn("w:ascii")) == "Arial"
                for name, value in (("b", "0"), ("bCs", "0"), ("i", "0"), ("iCs", "0"), ("u", "none"), ("color", "000000"), ("sz", "22"), ("szCs", "22")):
                    assert properties.find(qn(f"w:{name}")).get(qn("w:val")) == value
        record = logical_rows(output.tables[0])[0]
        hyperlink = record.header.cells[1]._tc.find(".//" + qn("w:hyperlink"))
        assert output.part.rels[hyperlink.get(qn("r:id"))].target_ref == "https://example.test/source"
        assert output.paragraphs[0].text == "Заголовок поза таблицею"
        assert output.tables[1].cell(0, 0).tables[0].cell(0, 0).text == "Вкладена таблиця"
        assert all(run.font.highlight_color is None for run in output.paragraphs[0].runs)

    def test_unselected_rows_lose_old_color_but_do_not_receive_new_highlights(self):
        raw = _table_docx([
            (["С. 1", "Текст"], ["С. 2", "Текст"]),
            (["С. 3", "Текст"], ["С. 4", "Текст"]),
        ])
        first = process_document(raw, HighlightOptions())
        second = process_document(first.document_bytes, HighlightOptions(first_row=2, last_row=2))
        document = Document(io.BytesIO(second.document_bytes))
        records = logical_rows(document.tables[0])
        first_colors = {r.get(qn("w:val")) for r in records[0].row._tr.iter(qn("w:highlight"))}
        second_colors = {r.get(qn("w:val")) for r in records[1].row._tr.iter(qn("w:highlight"))}
        assert first_colors == set()
        assert second_colors == {"yellow"}


class TestExactMarkerLayout:
    @pytest.mark.parametrize("font_name,font_size,width", [("Calibri", 14, 2.1), ("Arial", 16, 1.5), ("Times New Roman", 8, 3.0)])
    def test_wrapping_and_original_padding_cannot_offset_markers(self, font_name, font_size, width):
        data = _table_docx([(["", "", "c. 35", "Текст"], ["Широкі слова WWW " * 25, "https://example.test/" + "long/" * 35, "с. 88", "Текст"])])
        document = Document(io.BytesIO(data))
        row = document.tables[0].rows[0]
        for index, cell in enumerate(row.cells):
            cell.width = Inches(width + index * 0.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        for index, paragraph in enumerate((row.cells[0].paragraphs[2], row.cells[1].paragraphs[2])):
            paragraph.paragraph_format.space_before = Pt(20 * index + 10)
            paragraph.paragraph_format.line_spacing = 1.5 + index
            paragraph.paragraph_format.page_break_before = True
        raw = io.BytesIO()
        document.save(raw)
        result = process_document(raw.getvalue(), HighlightOptions(font_name=font_name, font_size=font_size))
        output = Document(io.BytesIO(result.document_bytes))
        record = logical_rows(output.tables[0])[0]
        assert result.stats.aligned_rows == 1
        assert _cell_texts(record.header.cells[0]) == ["", ""]
        assert [cell.paragraphs[0].text for cell in record.row.cells] == ["c. 35", "с. 88"]
        for cell in record.row.cells:
            assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            assert paragraph.paragraph_format.space_before == Pt(0)
            assert paragraph.paragraph_format.space_after == Pt(0)
            assert paragraph.paragraph_format.line_spacing == 1
            assert paragraph.paragraph_format.page_break_before is False
            assert paragraph.paragraph_format.keep_with_next is False
            assert cell._tc.tcPr.find(qn("w:tcMar")).find(qn("w:top")).get(qn("w:w")) == "0"
            assert cell._tc.tcPr.find(qn("w:tcBorders")).find(qn("w:top")).get(qn("w:val")) == "nil"

    def test_cells_without_optional_properties_can_be_aligned(self):
        document = Document(io.BytesIO(_table_docx([(["С. 1", "Текст"], ["Джерело", "С. 2", "Текст"])])))
        for cell in document.tables[0].rows[0].cells:
            cell._tc.remove(cell._tc.tcPr)
        raw = io.BytesIO()
        document.save(raw)
        result = process_document(raw.getvalue(), HighlightOptions())
        output = Document(io.BytesIO(result.document_bytes))
        record = logical_rows(output.tables[0])[0]
        assert [cell.paragraphs[0].text for cell in record.row.cells] == ["С. 1", "С. 2"]

    def test_second_run_keeps_logical_selection_text_blank_lines_and_hyperlinks(self):
        rows = [
            (["", "С. 1", "Текст", "", "Пояснення"], ["Джерело", "С. 2", "Текст"]),
            (["С. 3", "Текст"], ["Джерело 2", "С. 4", "Текст"]),
        ]
        data = _table_docx(rows, link_in_right=True)
        first = process_document(data, HighlightOptions())
        calls = []
        second = process_document(first.document_bytes, HighlightOptions(first_row=2, last_row=2, font_size=10), lambda *args: calls.append(args))
        document = Document(io.BytesIO(second.document_bytes))
        assert len(document.tables[0].rows) == 4
        records = logical_rows(document.tables[0])
        assert len(records) == 2
        assert calls == [(1, 1, 2)]
        assert second.stats.processed_rows == 1
        assert inspect_tables(data) == inspect_tables(second.document_bytes)
        assert _cell_texts(records[0].header.cells[0]) + _cell_texts(records[0].row.cells[0]) == rows[0][0]
        assert _has_hyperlink(records[0].header.cells[1])
        assert _cell_texts(records[1].row.cells[1]) == ["С. 4", "Текст"]

    def test_vertical_merge_is_not_restructured(self):
        document = Document(io.BytesIO(_table_docx([
            (["С. 1", "Текст"], ["Джерело", "С. 2", "Текст"]),
            ([""], [""]),
        ])))
        document.tables[0].cell(0, 0).merge(document.tables[0].cell(1, 0))
        raw = io.BytesIO()
        document.save(raw)
        result = process_document(raw.getvalue(), HighlightOptions())
        output = Document(io.BytesIO(result.document_bytes))
        assert len(output.tables[0].rows) == 2
        assert result.stats.processed_rows == 0
        assert result.stats.skipped_rows == 2
        assert len(result.warnings) == 2
