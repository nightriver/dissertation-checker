"""
Tests for the file-reading half of extractor.py: format dispatch, size guard,
PDF/DOCX line extraction and Word auto-numbering recovery.

Run: pytest tests/test_extractor_io.py
"""
import io

import pytest
from docx import Document

from parser import extractor
from parser.bibliography import parse_bibliography
from parser.extractor import (
    FileTooLargeError,
    ScannedPDFError,
    UnsupportedFormatError,
    extract_lines,
    extract_lines_from_docx,
    extract_lines_from_pdf,
)


class TestDispatch:
    def test_unknown_extension_raises(self):
        with pytest.raises(UnsupportedFormatError):
            extract_lines(b"data", "thesis.txt")

    def test_doc_is_rejected(self):
        # Legacy binary .doc is not readable by python-docx and must not be
        # mistaken for .docx.
        with pytest.raises(UnsupportedFormatError):
            extract_lines(b"data", "thesis.doc")

    def test_uppercase_extension_accepted(self, make_pdf):
        lines = extract_lines(make_pdf([["Hello world"]]), "THESIS.PDF")
        assert lines and lines[0]["page"] == 1

    def test_docx_dispatch(self, make_docx):
        lines = extract_lines(make_docx(["Перший абзац"]), "THESIS.DOCX")
        assert lines[0]["line"] == "Перший абзац"


class TestSizeGuard:
    def test_at_the_limit_is_accepted(self, monkeypatch, make_pdf):
        data = make_pdf([["Hello"]])
        monkeypatch.setattr(extractor, "MAX_FILE_SIZE", len(data))
        assert extract_lines_from_pdf(data)  # must not raise

    def test_one_byte_over_the_limit_raises(self, monkeypatch, make_pdf):
        data = make_pdf([["Hello"]])
        monkeypatch.setattr(extractor, "MAX_FILE_SIZE", len(data) - 1)
        with pytest.raises(FileTooLargeError):
            extract_lines_from_pdf(data)

    def test_guard_applies_to_docx_too(self, monkeypatch, make_docx):
        data = make_docx(["Текст"])
        monkeypatch.setattr(extractor, "MAX_FILE_SIZE", len(data) - 1)
        with pytest.raises(FileTooLargeError):
            extract_lines_from_docx(data)


class TestPdfExtraction:
    def test_pages_are_numbered_from_one(self, make_pdf):
        lines = extract_lines_from_pdf(make_pdf([["first"], ["second"]]))
        assert [item["page"] for item in lines] == [1, 2]

    def test_blank_lines_dropped(self, make_pdf):
        lines = extract_lines_from_pdf(make_pdf([["alpha", "", "beta"]]))
        assert [item["line"] for item in lines] == ["alpha", "beta"]

    def test_scanned_pdf_raises(self, empty_pdf):
        with pytest.raises(ScannedPDFError):
            extract_lines_from_pdf(empty_pdf)


class TestDocxExtraction:
    def test_empty_paragraphs_dropped(self, make_docx):
        lines = extract_lines_from_docx(make_docx(["Перший", "", "  ", "Другий"]))
        assert [item["line"] for item in lines] == ["Перший", "Другий"]

    def test_page_is_none(self, make_docx):
        lines = extract_lines_from_docx(make_docx(["Текст"]))
        assert lines[0]["page"] is None

    def test_trailing_whitespace_stripped(self, make_docx):
        lines = extract_lines_from_docx(make_docx(["Текст   "]))
        assert lines[0]["line"] == "Текст"


class TestWordAutoNumbering:
    """
    Word stores list numbers in numbering.xml, not in the paragraph text, so
    python-docx returns bibliography entries with no numbers at all. Without
    recovering them parse_bibliography finds zero entries and the whole
    analysis silently reports "every source is cited".
    """

    def test_numbers_are_recovered(self, make_numbered_docx):
        data = make_numbered_docx(
            ["СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ"],
            ["Абдуллаев Р.А. Полімери.", "Абрамович І. Відходи.", "Байцар Р.І. Якість."],
        )
        lines = extract_lines_from_docx(data)
        texts = [item["line"] for item in lines]

        assert texts[1].startswith("1. ")
        assert texts[2].startswith("2. ")
        assert texts[3].startswith("3. ")

    def test_recovered_numbers_make_the_list_parseable(self, make_numbered_docx):
        data = make_numbered_docx(
            ["СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ"],
            [f"Автор {i}. Назва праці номер {i}." for i in range(1, 21)],
        )
        lines = extract_lines_from_docx(data)
        parsed = parse_bibliography(lines[1:])

        assert len(parsed) == 20
        assert parsed[1].startswith("Автор 1.")
        assert parsed[20].startswith("Автор 20.")

    def test_plain_paragraphs_are_not_numbered(self, make_docx):
        lines = extract_lines_from_docx(make_docx(["Звичайний абзац без списку"]))
        assert lines[0]["line"] == "Звичайний абзац без списку"

    def test_headings_are_never_numbered(self):
        # A numbered Heading must keep its text intact, otherwise
        # _is_biblio_header stops matching "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ".
        document = Document()
        document.add_paragraph("СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", style="Heading 1")
        document.add_paragraph("Джерело перше.", style="List Number")
        buf = io.BytesIO()
        document.save(buf)

        lines = extract_lines_from_docx(buf.getvalue())
        assert lines[0]["line"] == "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ"
        assert lines[1]["line"].startswith("1. ")

    def test_bulleted_list_is_not_numbered(self):
        document = Document()
        document.add_paragraph("Перший пункт", style="List Bullet")
        document.add_paragraph("Другий пункт", style="List Bullet")
        buf = io.BytesIO()
        document.save(buf)

        lines = extract_lines_from_docx(buf.getvalue())
        assert [item["line"] for item in lines] == ["Перший пункт", "Другий пункт"]
