"""Підготовка, перевірка й безпечна обробка однієї таблиці DOCX."""

from __future__ import annotations

import io
from dataclasses import replace
from zipfile import BadZipFile, ZipFile, is_zipfile

from docx import Document

from table_highlighter.formatting import normalize_document
from table_highlighter.layout import align_page_markers, logical_rows, supports_alignment
from table_highlighter.matcher import align, left_statuses, right_statuses
from table_highlighter.types import (
    DocumentValidationError,
    HighlightOptions,
    HighlightResult,
    HighlightStats,
    RowWarning,
    TableSummary,
)
from table_highlighter.writer import (
    style_plain_paragraph,
    style_text_paragraph,
    supports_highlighting,
)
from table_highlighter.zones import annotate_left_cell, annotate_right_cell


MAX_DOCX_BYTES = 30 * 1024 * 1024


def _load_document(data: bytes) -> Document:
    if len(data) > MAX_DOCX_BYTES:
        raise DocumentValidationError("Файл більший за 30 МБ.")
    if not is_zipfile(io.BytesIO(data)):
        raise DocumentValidationError("Файл не є коректним контейнером DOCX.")
    try:
        with ZipFile(io.BytesIO(data)) as archive:
            if "word/document.xml" not in archive.namelist():
                raise DocumentValidationError("У контейнері немає основної частини документа Word.")
    except BadZipFile as error:
        raise DocumentValidationError("Файл не є коректним контейнером DOCX.") from error
    try:
        return Document(io.BytesIO(data))
    except Exception as error:
        raise DocumentValidationError("Не вдалося прочитати документ Word.") from error


def inspect_tables(data: bytes) -> tuple[TableSummary, ...]:
    """Повертає доступні таблиці без будь-яких змін документа."""
    document = _load_document(data)
    return tuple(
        TableSummary(index=index, row_count=len(logical_rows(table)), two_column_rows=sum(len(item.row.cells) == 2 for item in logical_rows(table)))
        for index, table in enumerate(document.tables)
    )


def _row_range(total: int, options: HighlightOptions) -> range:
    first = max(1, options.first_row)
    last = min(total, options.last_row if options.last_row is not None else total)
    return range(first - 1, last) if first <= last else range(0)


def _warning(row_number: int, message: str) -> RowWarning:
    return RowWarning(row_number=row_number, message=message)


def process_document(data: bytes, options: HighlightOptions, progress_callback=None) -> HighlightResult:
    """Підсвічує одну обрану таблицю й повертає новий DOCX без зміни входу."""
    if not 60 <= options.threshold <= 100:
        raise ValueError("Поріг схожості має бути від 60 до 100.")
    if not 8 <= options.font_size <= 16:
        raise ValueError("Розмір шрифту має бути від 8 до 16 pt.")
    document = _load_document(data)
    if not document.tables:
        raise DocumentValidationError("Документ не містить таблиць.")
    if not 0 <= options.table_index < len(document.tables):
        raise DocumentValidationError("Обраної таблиці не існує в документі.")

    table = document.tables[options.table_index]
    normalize_document(document, options.font_name, options.font_size)
    warnings: list[RowWarning] = []
    stats = HighlightStats()
    records = logical_rows(table)
    rows = tuple(_row_range(len(records), options))
    for position, row_index in enumerate(rows, 1):
        logical = records[row_index]
        row = logical.row
        number = row_index + 1
        if progress_callback:
            progress_callback(position, len(rows), number)
        if len(row.cells) != 2 or row.cells[0]._tc is row.cells[1]._tc:
            warnings.append(_warning(number, "Рядок не має двох окремих комірок."))
            stats = replace(stats, skipped_rows=stats.skipped_rows + 1)
            continue
        if not supports_alignment(row):
            warnings.append(_warning(number, "Рядок містить злиті комірки; підсвічування та вирівнювання пропущено."))
            stats = replace(stats, skipped_rows=stats.skipped_rows + 1)
            continue
        left_cell, right_cell = row.cells
        if not left_cell.text.strip() and not right_cell.text.strip():
            stats = replace(stats, skipped_rows=stats.skipped_rows + 1)
            continue
        left = annotate_left_cell(left_cell)
        right = annotate_right_cell(right_cell)
        if left.warning or right.warning:
            for message in (left.warning, right.warning):
                if message:
                    warnings.append(_warning(number, message))
            stats = replace(stats, skipped_rows=stats.skipped_rows + 1)
            continue
        if not left.text_paragraphs or not right.text_paragraphs:
            warnings.append(_warning(number, "Не знайдено порівнюваного тексту після маркерів сторінок."))
            stats = replace(stats, skipped_rows=stats.skipped_rows + 1)
            continue
        unsupported = [
            item for item in (*left.text_paragraphs, *right.text_paragraphs)
            if not supports_highlighting(item.paragraph)
        ]
        if unsupported:
            warnings.append(_warning(number, "Порівнювана зона містить непідтримуваний об'єкт Word; текст збережено, підсвічування та вирівнювання пропущено."))
            stats = replace(stats, skipped_rows=stats.skipped_rows + 1)
            continue

        alignment = align(left.text_paragraphs, right.text_paragraphs, options.threshold, options.relax_short_words)
        for item in left.paragraphs:
            if item.zone == "text":
                style_text_paragraph(item.paragraph, left_statuses(alignment)[item.index], options.font_name, options.font_size)
            else:
                style_plain_paragraph(item.paragraph, options.font_name, options.font_size)
        for item in right.paragraphs:
            if item.zone == "text":
                style_text_paragraph(item.paragraph, right_statuses(alignment)[item.index], options.font_name, options.font_size)
            else:
                style_plain_paragraph(item.paragraph, options.font_name, options.font_size)
        align_page_markers(logical, left, right, options.font_name, options.font_size)
        stats = replace(
            stats,
            processed_rows=stats.processed_rows + 1,
            exact_words=stats.exact_words + alignment.exact_words,
            fuzzy_words=stats.fuzzy_words + alignment.fuzzy_words,
            aligned_rows=stats.aligned_rows + 1,
        )

    output = io.BytesIO()
    document.save(output)
    return HighlightResult(output.getvalue(), stats, tuple(warnings))
