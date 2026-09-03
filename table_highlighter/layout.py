"""Вирівнювання маркерів спільною межею рядка, без оцінки переносів."""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from uuid import uuid4

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Row

from table_highlighter.formatting import plain_run_properties


_BOOKMARK_PREFIX = "_THA_"
_PROPERTY_ORDER = {
    qn("w:tcPr"): ("cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark", "headers", "cellIns", "cellDel", "cellMerge", "tcPrChange"),
    qn("w:tcBorders"): ("top", "start", "left", "bottom", "end", "right", "insideH", "insideV", "tl2br", "tr2bl"),
    qn("w:tcMar"): ("top", "start", "left", "bottom", "end", "right"),
}


@dataclass(frozen=True)
class LogicalRow:
    """Один запис порівняння: службова шапка та власне фрагменти."""

    row: _Row
    header: _Row | None = None


def _pair_names(row, suffix: str) -> set[str]:
    return {
        name[:-2]
        for element in row._tr.iter(qn("w:bookmarkStart"))
        if (name := element.get(qn("w:name"), "")).startswith(_BOOKMARK_PREFIX)
        and name.endswith(suffix)
    }


def logical_rows(table) -> tuple[LogicalRow, ...]:
    """Службові шапки не змінюють нумерацію записів при повторному запуску."""
    physical = tuple(table.rows)
    result = []
    index = 0
    while index < len(physical):
        if index + 1 < len(physical) and (
            _pair_names(physical[index], "_H") & _pair_names(physical[index + 1], "_B")
        ):
            result.append(LogicalRow(physical[index + 1], physical[index]))
            index += 2
        else:
            result.append(LogicalRow(physical[index]))
            index += 1
    return tuple(result)


def supports_alignment(row) -> bool:
    """Вертикальне злиття потребує окремої схеми і тут не перебудовується."""
    return len(row._tr.tc_lst) == 2 and all(
        not cell.xpath("./w:tcPr/w:vMerge | ./w:tcPr/w:hMerge | ./w:tcPr/w:gridSpan")
        for cell in row._tr.tc_lst
    )


def _property(parent, name: str, **attributes):
    element = parent.find(qn(f"w:{name}"))
    if element is None:
        element = OxmlElement(f"w:{name}")
        if parent.tag == qn("w:pPr"):
            parent.insert_element_before(element, "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange")
        else:
            order = _PROPERTY_ORDER[parent.tag]
            following = {qn(f"w:{key}") for key in order[order.index(name) + 1:]}
            position = next((index for index, child in enumerate(parent) if child.tag in following), len(parent))
            parent.insert(position, element)
    for key, value in attributes.items():
        element.set(qn(f"w:{key}"), str(value))
    return element


def _bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(start)
    paragraph._p.append(end)


def _mark_pair(header, body) -> None:
    existing_ids = [
        int(element.get(qn("w:id")))
        for element in body._tr.getroottree().getroot().iter(qn("w:bookmarkStart"))
        if element.get(qn("w:id"), "").isdigit()
    ]
    first_id = max(existing_ids, default=-1) + 1
    name = _BOOKMARK_PREFIX + uuid4().hex
    _bookmark(header.cells[0].paragraphs[0], name + "_H", first_id)
    _bookmark(body.cells[0].paragraphs[0], name + "_B", first_id + 1)


def _reset_row_height(row) -> None:
    properties = row._tr.trPr
    if properties is not None:
        for name in ("trHeight", "cantSplit", "tblHeader"):
            for element in tuple(properties.findall(qn(f"w:{name}"))):
                properties.remove(element)


def align_page_markers(logical: LogicalRow, left_zones, right_zones, font_name: str, font_size: int) -> None:
    """Зберігає всі абзаци, а маркери ставить на спільний верхній край."""
    row, header = logical.row, logical.header
    cells = tuple(row.cells)
    markers = [zones.paragraphs[zones.marker_index].paragraph for zones in (left_zones, right_zones)]
    prefixes = []
    for cell, marker in zip(cells, markers):
        children = list(cell._tc)
        prefixes.append([child for child in children[:children.index(marker._p)] if child.tag != qn("w:tcPr")])

    new_header = header is None and any(prefixes)
    if new_header:
        header_element = OxmlElement("w:tr")
        if row._tr.trPr is not None:
            header_element.append(deepcopy(row._tr.trPr))
        for cell in cells:
            header_cell = OxmlElement("w:tc")
            if cell._tc.tcPr is not None:
                header_cell.append(deepcopy(cell._tc.tcPr))
            header_element.append(header_cell)
        row._tr.addprevious(header_element)
        header = _Row(header_element, row._parent)

    if header is not None:
        for cell, prefix in zip(header.cells, prefixes):
            for child in prefix:
                cell._tc.append(child)
            if not len(cell._tc) or cell._tc[-1].tag != qn("w:p"):
                cell._tc.append(OxmlElement("w:p"))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            properties = cell._tc.get_or_add_tcPr()
            _property(_property(properties, "tcMar"), "bottom", w=0, type="dxa")
            _property(_property(properties, "tcBorders"), "bottom", val="nil")
            for paragraph in cell.paragraphs:
                # У Word keepNext тут притягує весь наступний рядок і лишає
                # великі порожні ділянки перед довгими фрагментами.
                paragraph.paragraph_format.keep_with_next = False
                # Новий порожній абзац теж має той самий кегль.
                ppr = paragraph._p.get_or_add_pPr()
                if ppr.find(qn("w:rPr")) is None:
                    ppr.insert_element_before(plain_run_properties(font_name, font_size), "w:sectPr", "w:pPrChange")
        _reset_row_height(header)
        if new_header:
            _mark_pair(header, row)

    _reset_row_height(row)
    for cell, marker in zip(cells, markers):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        properties = cell._tc.get_or_add_tcPr()
        _property(_property(properties, "tcMar"), "top", w=0, type="dxa")
        if header is not None:
            _property(_property(properties, "tcBorders"), "top", val="nil")
        # Успадковані відступи/розриви можуть розсунути навіть маркери в одному рядку.
        formatting = marker.paragraph_format
        formatting.space_before = Pt(0)
        formatting.space_after = Pt(0)
        formatting.line_spacing = 1
        formatting.left_indent = Pt(0)
        formatting.right_indent = Pt(0)
        formatting.first_line_indent = Pt(0)
        formatting.page_break_before = False
        formatting.keep_with_next = False
        formatting.keep_together = True
        _property(marker._p.get_or_add_pPr(), "snapToGrid", val=0)
