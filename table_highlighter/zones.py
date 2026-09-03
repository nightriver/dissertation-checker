"""Виділення службових і порівнюваних зон у комірках таблиці DOCX."""

from __future__ import annotations

import re
from dataclasses import dataclass

from docx.oxml.ns import qn


PAGE_MARKER_RE = re.compile(r"^\s*[СсCc]\.\s*\d+[\d\s\-–—]*\s*$")
COMMENT_RE = re.compile(r"^(?:\d+\.|Посилання на:|Там само$)")


@dataclass(frozen=True)
class ParagraphZone:
    """Абзац комірки та його роль у поточному рядку."""

    paragraph: object
    index: int
    zone: str


@dataclass(frozen=True)
class CellZones:
    """Повна розмітка комірки; усі абзаци присутні в результаті."""

    paragraphs: tuple[ParagraphZone, ...]
    marker_index: int | None
    warning: str | None = None

    @property
    def text_paragraphs(self) -> tuple[ParagraphZone, ...]:
        return tuple(item for item in self.paragraphs if item.zone == "text")


def paragraph_text(paragraph) -> str:
    """Повертає видимий текст, включно з текстом усередині гіперпосилань."""
    parts: list[str] = []
    for element in paragraph._p.iter():
        if element.tag == qn("w:t"):
            parts.append(element.text or "")
        elif element.tag == qn("w:tab"):
            parts.append("\t")
        elif element.tag in {qn("w:br"), qn("w:cr")}:
            parts.append("\n")
    return "".join(parts)


def is_page_marker(text: str) -> bool:
    """Розпізнає маркер сторінки порівнюваного фрагмента."""
    return bool(PAGE_MARKER_RE.match(text.strip()) or text.strip() == "--")


def annotate_left_cell(cell) -> CellZones:
    """Текст ліворуч розташований після першого маркера до першого пробілу."""
    paragraphs = tuple(cell.paragraphs)
    marker_index = next(
        (index for index, paragraph in enumerate(paragraphs) if is_page_marker(paragraph_text(paragraph))),
        None,
    )
    if marker_index is None:
        return CellZones(
            tuple(ParagraphZone(paragraph, index, "plain") for index, paragraph in enumerate(paragraphs)),
            None,
            "Не знайдено маркер сторінки в лівій комірці.",
        )

    text_open = True
    result: list[ParagraphZone] = []
    for index, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        if index <= marker_index:
            zone = "plain"
        elif text_open and not text.strip():
            text_open = False
            zone = "plain"
        else:
            zone = "text" if text_open else "plain"
        result.append(ParagraphZone(paragraph, index, zone))
    return CellZones(tuple(result), marker_index)


def annotate_right_cell(cell) -> CellZones:
    """Текст праворуч іде після останнього маркера до коментаря."""
    paragraphs = tuple(cell.paragraphs)
    marker_index: int | None = None
    for index, paragraph in enumerate(paragraphs):
        if is_page_marker(paragraph_text(paragraph)):
            marker_index = index
    if marker_index is None:
        return CellZones(
            tuple(ParagraphZone(paragraph, index, "plain") for index, paragraph in enumerate(paragraphs)),
            None,
            "Не знайдено маркер сторінки в правій комірці.",
        )

    text_open = True
    result: list[ParagraphZone] = []
    for index, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        if index <= marker_index:
            zone = "plain"
        elif text_open and COMMENT_RE.match(text.strip()):
            text_open = False
            zone = "plain"
        else:
            zone = "text" if text_open else "plain"
        result.append(ParagraphZone(paragraph, index, zone))
    return CellZones(tuple(result), marker_index)
